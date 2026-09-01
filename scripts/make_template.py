import shutil
from docx import Document
from pathlib import Path

BASE_DIR = Path(".")
src_path = BASE_DIR / "templates" / "Darnell.docx"
backup_path = BASE_DIR / "templates" / "Darnell_backup.docx"

doc = Document(backup_path)

# 1. Update Cover Paragraphs
for p in doc.paragraphs:
    if "TAHUN ANGGARAN" in p.text:
        p.text = "TAHUN ANGGARAN {{ report_year }}"
    elif p.text.strip() == "Juli 2026":
        p.text = "{{ report_month_upper }} {{ report_year }}"
    elif "Pada hari ini" in p.text:
        p.text = "Pada hari ini {{ bast_day_name }} tanggal {{ bast_date_terbilang }} bulan {{ bast_month_name }} tahun {{ bast_year_terbilang }} ({{ bast_date_numeric }}), kami yang bertandatangan dibawah ini :"
    elif p.text.startswith("Nama") and ":" in p.text:
        p.text = "Nama\t\t\t: {{ full_name }}"
    elif p.text.startswith("NIK") and ":" in p.text:
        p.text = "NIK\t\t\t: {{ nik }}"
    elif p.text.startswith("Penempatan") and ":" in p.text:
        p.text = "Penempatan\t\t: {{ placement }}"
    elif p.text.startswith("Area Pekerjaan") and ":" in p.text:
        p.text = "Area Pekerjaan\t: {{ area }}"
    elif "progres kemajuan pekerjaan pada Bulan" in p.text:
        p.text = "Dengan ini kami sampaikan bahwa progres kemajuan pekerjaan pada Bulan {{ report_month_upper }} {{ report_year }} sampai saat ini telah mencapai 100% sesuai dengan syarat-syarat dan kontrak."

# 2. Update Table 0 (Cover metadata)
t0 = doc.tables[0]
t0.rows[0].cells[2].text = "{{ name_upper }}"
t0.rows[1].cells[2].text = "{{ area }}"
t0.rows[2].cells[2].text = "{{ placement }}"
t0.rows[3].cells[2].text = "{{ report_month_upper }} {{ report_year }}"

# 3. Update Table 1 (Signature table)
t1 = doc.tables[1]
t1.rows[0].cells[0].text = t1.rows[0].cells[0].text.replace("Darnell Ignasius, S. Kom.", "{{ full_name }}")

# 4. Update Table 2 (Activity Rows Loop using 3-row pattern)
t2 = doc.tables[2]
# Remove all rows except header (row 0) and template data (row 1)
while len(t2.rows) > 2:
    row_to_remove = t2.rows[-1]._tr
    t2._tbl.remove(row_to_remove)

# Data row is row 1
t2.rows[1].cells[0].text = "{{ r.date_label }}"
t2.rows[1].cells[1].text = "{{ r.activity }}"
t2.rows[1].cells[2].text = "{{ r.category }}"

# Insert {%tr for r in rows %} before row 1
tr_for_row = t2.add_row()
tr_for_row.cells[0].text = "{%tr for r in rows %}"
# Move it before row 1 in XML
t2._tbl.insert(t2._tbl.index(t2.rows[1]._tr), tr_for_row._tr)

# Insert {%tr endfor %} after data row
tr_end_row = t2.add_row()
tr_end_row.cells[0].text = "{%tr endfor %}"

# 5. Remove Table 3 (Static Images)
if len(doc.tables) > 3:
    t3 = doc.tables[3]
    t3._tbl.getparent().remove(t3._tbl)

# Clean up empty paragraphs around LAMPIRAN and format heading
for idx, p in enumerate(list(doc.paragraphs)):
    if p.text.strip() == "LAMPIRAN":
        # Remove following empty paragraphs containing old page breaks
        for next_p in list(doc.paragraphs)[idx+1:]:
            if not next_p.text.strip():
                p_elem = next_p._p
                p_elem.getparent().remove(p_elem)
            else:
                break
        break

# Add Jinja images loop directly after LAMPIRAN
doc.add_paragraph("{%p for img in images %}")
doc.add_paragraph("{{ img }}")
doc.add_paragraph("{%p endfor %}")

doc.save(src_path)
print(f"Successfully generated Jinja2 template at {src_path}")
